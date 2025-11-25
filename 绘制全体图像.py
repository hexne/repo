from docx import Document
import matplotlib.pyplot as plt

# 设置中文字体支持
plt.rcParams['font.sans-serif'] = ['SimHei']  # 或者 ['Microsoft YaHei']
plt.rcParams['axes.unicode_minus'] = False

# 读取 Word 文档
doc_path = r'C:\Users\hexne\Desktop\改进统计.docx'
doc = Document(doc_path)

# 提取段落文本作为标题
paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
tables = doc.tables

# 指标字段及对应列索引（从0开始）
metrics = [
    ("准确率", 1),
    ("召回率", 2),
    ("mAP50", 3),
    ("mAP95", 4),
    ("F1", 5),
    ("时间(s)", 6)
]

para_index = 0

for metric_name, col_index in metrics:
    raw_data = []

    para_index = 0  # 每次从头开始匹配标题
    for table in tables:
        # 获取表格前的段落作为标题
        while para_index < len(paragraphs):
            title = paragraphs[para_index]
            para_index += 1
            if title:
                break

        # 遍历表格行（跳过表头）
        for i, row in enumerate(table.rows[1:], start=1):
            cells = row.cells
            if cells[0].text.strip() != "5":  # 只保留第一列为 "5e" 的行
                continue
            try:
                value = float(cells[col_index].text.strip())
                label = f"{title}{i}"
                raw_data.append((label, value))
            except:
                continue

    # 按指标值降序排序
    sorted_data = sorted(raw_data, key=lambda x: x[1], reverse=True)
    labels = [item[0] for item in sorted_data]
    values = [item[1] for item in sorted_data]

    # 绘制柱状图
    plt.figure(figsize=(max(10, len(labels) * 0.5), 6))
    bars = plt.bar(labels, values, color='skyblue')

    # 添加数值标签
    for bar in bars:
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width() / 2, height, f'{height:.4f}',
                 ha='center', va='bottom', fontsize=8)

    plt.xticks(rotation=90)
    plt.ylabel(metric_name)
    plt.title(f'{metric_name} 柱状图（按数值降序）')
    plt.tight_layout()
    plt.show()
